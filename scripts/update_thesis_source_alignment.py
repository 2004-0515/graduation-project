from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


SOURCE = Path(
    r"C:\Users\Administrator\Desktop\我的毕业设计材料\毕业设计终极版\初稿终极版\真终极版压缩包\终极修复版\2203010614张津滔毕业设计定稿.docx"
)
OUT_DIR = Path(r"D:\graduation-project\docs\thesis-alignment")
OUTPUT = OUT_DIR / "2203010614张津滔毕业设计定稿-源码统一修订版.docx"


def iter_blocks(doc: Document):
    body = doc.element.body
    para_index = 0
    table_index = 0
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc), "p", para_index
            para_index += 1
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc), "t", table_index
            table_index += 1


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    style = paragraph.style
    alignment = paragraph.alignment
    paragraph.clear()
    paragraph.style = style
    paragraph.alignment = alignment
    paragraph.add_run(text)


def normalize_text(text: str) -> str:
    return text.replace("\xa0", " ").strip()


def find_caption_table(doc: Document, caption: str) -> Table:
    seen_caption = False
    for block, kind, _ in iter_blocks(doc):
        if kind == "p" and caption in normalize_text(block.text):
            seen_caption = True
            continue
        if seen_caption and kind == "t":
            return block
        if seen_caption and kind == "p" and normalize_text(block.text):
            raise RuntimeError(f"Caption found but next table missing: {caption}")
    raise RuntimeError(f"Caption not found: {caption}")


def set_table_rows(table: Table, rows: list[list[str]]) -> None:
    if not rows:
        return

    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)

    for r_idx, row in enumerate(rows):
        while len(table.rows[r_idx].cells) < len(row):
            raise RuntimeError("Target table has fewer columns than expected")
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = value


def insert_paragraph_after(doc: Document, after_element, text: str, style_name: str):
    paragraph = doc.add_paragraph(text)
    try:
        paragraph.style = style_name
    except KeyError:
        pass
    after_element.addnext(paragraph._p)
    return paragraph._p


def insert_table_after(doc: Document, after_element, rows: list[list[str]], style):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    if style is not None:
        table.style = style
    set_table_rows(table, rows)
    after_element.addnext(table._tbl)
    return table._tbl


def replace_paragraphs(doc: Document) -> None:
    replacements = {
        "关键词：理性消费；购物商城系统；前后端分离；Springboot；Vue3":
            "关键词：理性消费；购物商城系统；前后端分离；SpringBoot；Vue3",
        "（2）清单管理策略与延迟回报系统的结合。参考了行为经济学里“冷静期”那个概念，在系统里加了个心愿清单。用户把商品丢进去之后，可以自己选时间的冻结期，这段时间里只能看不能买，算是强制冷静一下。等这个期限过了，商品自动恢复正常，能正常下单。说白了就是想给用户在付款之前多留点考虑时间，冲动消费的情况应该能少点。另外还做了个价格追踪，商品上架时先记一笔价格，后面要是降价或者涨价了，系统会给用户弹个提醒，省得他们错过变动。用户可以将感兴趣的商品添加至购物车或心愿清单，以便后续对比和决策。":
            "（2）清单管理策略与延迟回报系统的结合。系统在商品详情页提供心愿清单入口，用户可以为暂时拿不定的商品设置冷静期，并填写加入原因。冷静期内，该商品主要在心愿清单中保留和提醒；用户进入心愿清单时，系统会刷新已经到期的记录状态，到期后用户可从清单页继续前往商品详情页购买。价格追踪功能会记录商品价格变化，并在满足降价提醒条件时向用户提示，帮助用户在购买前获得更充分的信息。",
        "系统架构中，管理员拥有最高权限，负责系统管理和维护。管理后台需涵盖用户管理、商品管理、订单管理、数据统计及系统设置等核心功能模块，具体包括用户列表与详情查看、商品审核与上下架、订单状态跟踪、交易数据分析，以及权限与角色的动态分配。商品和文件由管理员直接发布或者上传后，系统会自动审核，不需要人工审核。这些操作都会被系统记录，保证平台正常运转以及可追溯性。通过上述明确的角色与权限划分，系统实现了前台业务与后台管理的有效分离，保障多角色协同下的电商业务高效与安全运转。":
            "系统架构中，管理员拥有最高权限，负责系统管理和维护。管理后台涵盖用户管理、商品管理、订单管理、数据统计及系统设置等核心功能模块，具体包括用户列表与详情查看、商品审核与上下架、订单状态跟踪、交易数据分析，以及权限与角色管理。卖家提交商品后，系统将其置为待审核状态；头像、商品图片等受管控文件上传后，也需要管理员在后台进行人工审核。审核通过后相关内容才正式生效，审核拒绝时系统记录原因并反馈给提交者。通过上述角色与权限划分，系统实现了前台业务与后台管理的有效分离，保障多角色协同下的电商业务正常运行。",
        "6.文件审阅：对用户以及商家上传图片进行检查，看是否符合规定。":
            "6.商品审核：对商家提交的商品信息、图片和价格变更进行审核，判断其是否符合平台上架要求。",
        "用户对暂时拿不定的商品，可以先加入心愿清单，并设置一段等待时间。等待期结束后，系统会向用户发送相应提示。商品也可单独设置观察期（默认7天），在此期间系统跟踪商品变化情况，并在观察期结束后给出提醒，帮助用户减少冲动下单行为。如表3.2所示。":
            "用户对暂时拿不定的商品，可以先加入心愿清单，并设置一段冷静期。系统页面提供1天、3天、7天和14天等选项，默认值为3天。用户查看心愿清单时，系统会根据冷静期结束时间刷新记录状态；到期后，用户可以从心愿清单继续前往商品详情页购买，从而减少冲动下单行为。如表3.2所示。",
        "（2）心愿单管理活动图可以用来表示系统的特色功能的主要业务流程。用户在商品详情页点击加入心愿单，系统显示冷静期设置对话框，用户选择冷静期天数（默认7天）。系统验证该商品是否已存在于用户心愿单中，验证通过后生成心愿记录，状态为冷静期，并计算冷静期结束时间。系统通过定时任务每日轮询心愿单冷静期状态，届满后自动将状态更新为可购买，并向用户发送提醒。用户查看心愿单时，若冷静期已过，可将商品加入购物车或从清单中移除。如图4.3所示。":
            "（2）心愿单管理活动图可以用来表示系统的特色功能的主要业务流程。用户在商品详情页点击加入心愿单，系统显示冷静期设置对话框，用户选择冷静期天数（默认3天）。系统验证该商品是否已存在于用户心愿单中，验证通过后生成心愿记录，状态为冷静中，并计算冷静期结束时间。用户查看心愿清单时，系统会刷新已经到期的记录状态；若冷静期已过，可从清单页继续前往商品详情页购买，也可以移除该商品。如图4.3所示。",
        "图4.8  心愿单管理时序图（2）心愿单管理时序图是利用了冷静期的概念来实现一个有时间缓冲的心愿单。当消费者把某件商品加入到心愿单中，系统会进行检查，如果这个商品还没有被加入到数据库中，那么就会生成一条状态是冷静期的信息并且设置一个七天的冷静期结束时间。系统通过定时任务每日轮询心愿单冷静期状态。根据不同的冷静期情况，界面上会有不同的操作：如果冷静期已经过了，那么商品的状态就会变成可以购买并且可以点击下单。如果冷静期还没有过，那么就会显示还剩多少时间才能购买。如图4.8所示。":
            "图4.8  心愿单管理时序图（2）心愿单管理时序图利用冷静期机制实现带有时间缓冲的心愿清单。用户在商品详情页提交商品ID、冷静期天数和加入原因后，系统先检查同一用户是否已有该商品的有效心愿记录；若不存在，则生成一条冷静中记录并计算冷静期结束时间。用户进入心愿清单页面时，服务层会刷新已到期记录的状态。若记录已到期，页面允许用户继续前往商品详情页购买；若仍在冷静期内，则显示剩余时间和状态提示。如图4.8所示。",
        "（5）订单取消活动图覆盖了用户提交申请后的系统处理链路。用户先查看订单明细，再发起取消请求；系统收到后核验当前状态，依所处阶段分流处理：待支付订单直接关闭，库存自动回滚；已付款但尚未发货的订单，用户需填写取消原因并提交申请，系统记录申请后将其置为“取消申请中”，并向管理员推送待审核通知，管理员审核通过后触发退款并同步恢复库存，审核不通过则订单改回待发货状态并通知用户；已进入履约准备阶段（如已分配仓库或打印面单）的订单，同样进入管理员审核流程；已发货或已完成的订单，系统直接驳回请求。如图4.6所示。":
            "（5）订单取消活动图覆盖了用户发起取消后的系统处理链路。用户先查看订单明细，再发起取消请求；系统收到后核验当前状态并分流处理：待支付订单可直接关闭并恢复库存；已支付且处于待发货状态的订单，用户需提交取消申请，系统将其置为“申请取消中”，并交由管理员审核。管理员审核通过后，订单改为已取消并同步恢复库存；审核不通过时，订单恢复为待发货状态。已发货、已完成等状态的订单不再进入取消流程，系统直接提示当前状态不允许取消。如图4.6所示。",
        "图4.11  订单取消时序图（5）订单取消时序图主要表达系统如何按订单阶段差异化处理取消请求。用户提交取消申请，系统先拉取订单当前状态，随后分流：待支付订单直接关闭并回滚库存；待发货订单，系统记录申请后改状态为“取消申请中”，并向管理员发送待审核通知，管理员审核通过后执行退款与库存恢复，若审核不通过则改回待发货并通知用户；已发货或已完成的订单，系统拒绝并告知用户该阶段不可取消。时序图上，待支付分支用虚线返回，履约阶段用实线转管理员，视觉区分度更明显。如图4.11所示。":
            "图4.11  订单取消时序图（5）订单取消时序图主要表达系统如何按订单阶段差异化处理取消请求。用户提交取消操作后，系统先读取订单当前状态并进行判断：待支付订单直接取消并恢复库存；待发货订单需要先提交取消申请，系统将订单状态更新为“申请取消中”，随后由管理员在后台审核。审核通过后订单变为已取消并恢复库存，审核拒绝后订单恢复为待发货状态；已发货或已完成的订单，系统直接返回不可取消提示。如图4.11所示。",
        "（4）订单表是主要的数据存储，系统性地保存了用户的交易信息以及变化过程，展示了订单管理系统的主要字段，包括订单唯一标识符、订单编号、用户ID、订单总金额、实付金额、优惠金额、订单状态（待支付、待发货、待签收、已完结、已取消等）、支付状态、支付方式、收货人姓名、联系电话、收货地址、订单备注、下单时间、付款时间、发货时间和订单结束时间。此数据模型可以实现对一个订单从产生到结束整个过程进行跟踪管理，如表4.4所示。":
            "（4）订单表是交易流程中的核心数据表，用于保存用户下单后的订单编号、用户ID、订单总金额、实付金额、支付方式、支付状态和订单状态等信息。订单状态包括待支付、待发货、待收货、已完成、已取消、退款中和申请取消中等取值；收货信息以地址快照形式保存，便于后续回看。该表还记录优惠券抵扣、下单时间、支付时间、发货时间和完成时间，用于跟踪订单从创建到结束的全过程，如表4.4所示。",
        "（9）心愿清单表是用于保存用户所选商品信息的数据表，同时包含帮助用户理性消费的冷静期控制功能。该表包括心愿单ID、用户ID、商品ID、冷静期天数、添加原因、状态、可购买时间和创建时间等字段。用户将商品加入心愿清单后，系统会根据冷静期天数计算可购买时间，在冷静期内主要用于提醒和暂存，冷静期结束后再允许用户继续购买，从而减少冲动消费行为，同时也方便用户回顾购买动机并管理后续购买计划，如表4.9所示。":
            "（9）心愿清单表用于保存用户暂时不立即购买的商品信息，同时承担冷静期控制功能。该表包括心愿单ID、用户ID、商品ID、加入时价格、冷静期天数、冷静期结束时间、状态、添加原因和创建时间等字段。用户将商品加入心愿清单后，系统会根据冷静期天数计算结束时间；用户查看清单时，系统刷新已到期记录的状态，冷静期结束后再允许从清单页继续前往购买，从而减少冲动消费行为，如表4.9所示。",
        "在商城的基础交易功能之外，系统还结合课题方向做了几个和理性消费相关的模块。用户可以按月设定消费预算，系统会在接近预算上限时给出提示；心愿单加入了冷静期机制，商品添加后并不能马上购买，需要等待设定的天数过后才能操作。另外还有价格提醒、消费月报和消费成就这几个功能，通过分析用户的消费行为，为用户提供个性化、有针对性的消费建议，从而促进理智消费。角色权限方面分了三种：普通用户可完成商品浏览、搜索、加入购物车等全流程消费行为；商家运营角色则拥有独立后台，支持商品上架下架、库存动态管理等核心经营能力；管理员则承担商品审核、取":
            "在商城的基础交易功能之外，系统还结合课题方向实现了理性消费相关模块。用户可以按月设定消费预算，系统会在接近预算上限时给出提示；心愿单加入冷静期机制，商品加入清单后会先进入暂缓决策状态，冷静期结束后可从清单页继续前往购买。另外，价格提醒、消费月报和消费成就等功能会基于用户的真实消费数据展示统计结果和提示信息。角色权限方面分为普通用户、商家和管理员：普通用户可完成商品浏览、搜索、加入购物车等消费行为；商家拥有商品维护和订单处理能力；管理员负责商品审核、订单管理和平台数据维护。",
        "订单流程设计了6条用例，覆盖提交订单、支付、取消审核、卖家发货、权限验证和确认收货。这条链路涉及买家、卖家和管理员三个角色，测试时需要切换账号。其中权限验证是用买家身份请求卖家发货接口，系统返回403，说明角色隔离生效。整个流程从下单到确认收货能按步骤走通，没有出现状态断掉的情况。测试结果均通过，。结果如表6-4所示。运行结果如图6.4所示。":
            "订单流程设计了6条用例，覆盖提交订单、支付、取消申请审核、卖家发货、权限验证和确认收货。这条链路涉及买家、卖家和管理员三个角色，测试时需要切换账号。其中权限验证是用买家身份请求卖家发货接口，系统返回403，说明角色隔离生效。整个流程从下单到确认收货能按步骤走通，没有出现状态断掉的情况。测试结果均通过，结果如表6-4所示。运行结果如图6.4所示。",
    }

    for paragraph in doc.paragraphs:
        text = normalize_text(paragraph.text)
        if text in replacements:
            set_paragraph_text(paragraph, replacements[text])
        elif "Springboot" in text:
            set_paragraph_text(paragraph, paragraph.text.replace("Springboot", "SpringBoot"))
        elif "取消审核中" in text or "取消申请中" in text:
            set_paragraph_text(
                paragraph,
                paragraph.text.replace("取消审核中", "申请取消中").replace("取消申请中", "申请取消中"),
            )
        elif "图4.8" in text and "心愿单管理时序图" in text:
            set_paragraph_text(
                paragraph,
                "图4.8  心愿单管理时序图（2）心愿单管理时序图利用冷静期机制实现带有时间缓冲的心愿清单。用户在商品详情页提交商品ID、冷静期天数和加入原因后，系统先检查同一用户是否已有该商品的有效心愿记录；若不存在，则生成一条冷静中记录并计算冷静期结束时间。用户进入心愿清单页面时，服务层会刷新已到期记录的状态。若记录已到期，页面允许用户继续前往商品详情页购买；若仍在冷静期内，则显示剩余时间和状态提示。如图4.8所示。",
            )
        elif "在商城的基础交易功能之外" in text and "商品添加后并不能马上购买" in text:
            set_paragraph_text(
                paragraph,
                "在商城的基础交易功能之外，系统还结合课题方向实现了理性消费相关模块。用户可以按月设定消费预算，系统会在接近预算上限时给出提示；心愿单加入冷静期机制，商品加入清单后会先进入暂缓决策状态，冷静期结束后可从清单页继续前往购买。另外，价格提醒、消费月报和消费成就等功能会基于用户的真实消费数据展示统计结果和提示信息。角色权限方面分为普通用户、商家和管理员：普通用户可完成商品浏览、搜索、加入购物车等消费行为；商家拥有商品维护和订单处理能力；管理员负责商品审核、订单管理和平台数据维护。整体来看，系统既保留了普通商城的交易能力，也把理性消费引导功能嵌入到了购物流程中。",
            )


def update_use_case_tables(doc: Document) -> None:
    table32 = find_caption_table(doc, "表3.2")
    set_table_rows(table32, [
        ["用例编号：", "UC002"],
        ["用例名称：", "心愿单管理"],
        ["简要描述：", "用户可以将感兴趣但暂时不确定是否购买的商品加入心愿清单，并设置冷静期；冷静期结束后，系统在用户查看清单时刷新状态，用户可继续前往商品详情页购买。"],
        ["参与者：", "用户"],
        ["前置条件：", "用户已经登录系统，正在查看商品详情页。"],
        ["后置条件：", "商品被加入用户心愿清单，记录为冷静中状态；冷静期结束后，用户可从心愿清单继续前往商品详情页购买或移除该记录。"],
        ["基本流程：", "1.用户在商品详情页点击加入心愿单按钮。2.系统显示冷静期设置对话框。3.用户选择冷静期天数，页面默认值为3天，并可填写加入原因。4.系统校验该商品是否已存在于用户有效心愿清单中。5.校验通过后，系统生成心愿单记录，状态为冷静中。6.系统根据冷静期天数计算并保存冷静期结束时间。7.系统显示“已加入心愿单”提示。8.用户进入心愿清单页面时，系统刷新已到期记录的状态。9.若冷静期结束，页面显示继续购买入口；若未结束，页面显示剩余冷静时间。"],
        ["备选流程：", "E-1：若当前用户的心愿清单中已包含该商品，系统提示“该商品已在想要清单中”并终止本次操作。E-2：冷静期届满后，用户点击“去购买”，系统先将该心愿记录标记为已购买，再跳转至商品详情页，由用户继续完成购买流程。E-3：用户不再需要该商品时，可从心愿清单中移除记录。"],
        ["业务规则：", "1.心愿单状态标识定义如下：0—冷静中，1—可购买，2—已购买，3—已移除。2.冷静期页面默认值为3日，支持用户选择1日、3日、7日和14日。3.同一用户对同一商品只保留一条冷静中或可购买状态的有效记录。4.系统在查询心愿清单时刷新已到期记录状态，不依赖后台定时任务。"],
    ])

    table33 = find_caption_table(doc, "表3.3")
    table33.cell(8, 1).text = (
        "1.审核状态标识：0表示待审核，1表示审核通过，2表示审核拒绝。"
        "2.商品状态：0表示下架，1表示上架。"
        "3.卖家提交商品后进入待审核状态，必须由管理员人工审核；管理员后台直接创建的商品可直接置为审核通过。"
        "4.审核拒绝时必须填写拒绝原因。"
        "5.商品被驳回后，卖家可以修改商品信息并重新提交平台审核。"
    )

    table35 = find_caption_table(doc, "表3.5")
    set_table_rows(table35, [
        ["用例编号：", "UC005"],
        ["用例名称：", "订单取消"],
        ["简要描述：", "系统根据订单所处生命周期对取消请求进行分流处理，待支付订单可直接取消，待发货订单需提交取消申请并由管理员审核。"],
        ["参与者：", "用户、管理员"],
        ["前置条件：", "用户已经登录系统，并且账户下存在可查看的订单记录。"],
        ["后置条件：", "符合条件的订单被取消并恢复对应商品库存；不符合取消条件的订单保持原状态，并向用户返回提示信息。"],
        ["基本流程：", "1.用户进入订单列表页面。2.系统显示该用户的订单信息。3.用户选择一笔待支付订单并执行取消操作。4.系统弹出确认提示。5.用户确认后，系统将订单状态改为已取消。6.系统恢复该订单占用的商品库存并回退销量。7.系统提示订单取消成功。"],
        ["备选流程：", "E-1：若订单处于待发货状态，用户提交取消申请后，系统将订单状态改为申请取消中，并交由管理员审核。管理员同意后，订单改为已取消并恢复库存；管理员拒绝后，订单恢复为待发货状态。E-2：当订单处于待收货、已完成等状态时，系统提示当前状态不允许取消，用例结束。"],
        ["业务规则：", "1.订单状态标识：0表示待支付，1表示待发货，2表示待收货，3表示已完成，4表示已取消，5表示退款中，6表示申请取消中。2.待支付订单可以直接取消，无需审核。3.待发货订单如需取消，必须先提交取消申请，由管理员审核通过后才能取消。4.待收货或已完成的订单不能直接取消。5.取消订单后释放商品库存并回退销量。"],
    ])


def update_physical_tables(doc: Document) -> None:
    header = ["字段名称", "类型", "长度", "字段说明", "主键", "默认值"]

    set_table_rows(find_caption_table(doc, "表4.4"), [
        header,
        ["id", "bigint", "20", "订单ID", "是", "-"],
        ["order_no", "varchar", "50", "订单编号", "否", "-"],
        ["user_id", "bigint", "20", "用户ID", "否", "-"],
        ["total_amount", "decimal", "10,2", "订单总金额", "否", "-"],
        ["pay_amount", "decimal", "10,2", "实付金额", "否", "NULL"],
        ["payment_method", "tinyint", "4", "支付方式", "否", "1"],
        ["payment_status", "tinyint", "4", "支付状态", "否", "0"],
        ["order_status", "tinyint", "4", "订单状态", "否", "0"],
        ["shipping_address", "text", "-", "收货地址快照", "否", "NULL"],
        ["payment_time", "datetime", "-", "支付时间", "否", "NULL"],
        ["shipping_time", "datetime", "-", "发货时间", "否", "NULL"],
        ["end_time", "datetime", "-", "完成时间", "否", "NULL"],
        ["remark", "varchar", "200", "订单备注", "否", "NULL"],
        ["coupon_id", "bigint", "20", "使用优惠券ID", "否", "NULL"],
        ["coupon_discount", "decimal", "10,2", "优惠券抵扣金额", "否", "NULL"],
        ["created_time", "datetime", "-", "创建时间", "否", "CURRENT_TIMESTAMP"],
        ["updated_time", "datetime", "-", "更新时间", "否", "CURRENT_TIMESTAMP"],
    ])

    set_table_rows(find_caption_table(doc, "表4.5"), [
        header,
        ["id", "bigint", "20", "商品ID", "是", "-"],
        ["name", "varchar", "100", "商品名称", "否", "-"],
        ["description", "text", "-", "商品描述", "否", "NULL"],
        ["category_id", "bigint", "20", "分类ID", "否", "-"],
        ["price", "decimal", "10,2", "商品价格", "否", "-"],
        ["original_price", "decimal", "10,2", "原价", "否", "NULL"],
        ["pending_price", "decimal", "10,2", "待审核价格", "否", "NULL"],
        ["pending_original_price", "decimal", "10,2", "待审核原价", "否", "NULL"],
        ["stock", "int", "11", "库存数量", "否", "0"],
        ["version", "bigint", "20", "乐观锁版本号", "否", "0"],
        ["sales", "int", "11", "销量", "否", "0"],
        ["status", "tinyint", "4", "商品状态", "否", "1"],
        ["main_image", "varchar", "200", "主图URL", "否", "NULL"],
        ["images", "text", "-", "商品图片列表", "否", "NULL"],
        ["seller_id", "bigint", "20", "卖家用户ID", "否", "-"],
        ["seller_name", "varchar", "50", "卖家用户名", "否", "-"],
        ["audit_status", "tinyint", "4", "审核状态", "否", "1"],
        ["audit_remark", "varchar", "200", "审核备注", "否", "NULL"],
        ["audit_time", "datetime", "-", "审核时间", "否", "NULL"],
        ["ad_video", "varchar", "500", "广告视频URL", "否", "NULL"],
        ["ad_video_duration", "int", "11", "广告时长（秒）", "否", "NULL"],
        ["ad_video_enabled", "tinyint", "4", "广告启用状态", "否", "0"],
        ["created_time", "datetime", "-", "创建时间", "否", "CURRENT_TIMESTAMP"],
        ["updated_time", "datetime", "-", "更新时间", "否", "CURRENT_TIMESTAMP"],
    ])

    set_table_rows(find_caption_table(doc, "表4.9"), [
        header,
        ["id", "bigint", "20", "心愿单ID", "是", "-"],
        ["user_id", "bigint", "20", "用户ID", "否", "-"],
        ["product_id", "bigint", "20", "商品ID", "否", "-"],
        ["added_price", "decimal", "10,2", "加入时价格", "否", "NULL"],
        ["cooling_days", "int", "11", "冷静期天数", "否", "3"],
        ["cooling_end_time", "datetime", "-", "冷静期结束时间", "否", "NULL"],
        ["status", "tinyint", "4", "状态", "否", "0"],
        ["reason", "varchar", "500", "添加原因", "否", "NULL"],
        ["created_time", "datetime", "-", "创建时间", "否", "CURRENT_TIMESTAMP"],
    ])


def append_new_physical_tables(doc: Document) -> None:
    table411 = find_caption_table(doc, "表4.11")
    anchor = table411._tbl
    style_name = "Body Text"
    caption_style = "Caption"
    table_style = table411.style
    header = ["字段名称", "类型", "长度", "字段说明", "主键", "默认值"]

    blocks = [
        ("p", "（12）价格历史表用于记录商品价格随时间发生的变化，是商品详情页展示价格走势、最低价、最高价和平均价格的主要数据来源。该表包括价格历史ID、商品ID、记录价格、原价、记录时间、变化类型、变化金额和变化比例等字段。商品新增或价格调整时，系统会写入价格历史记录，便于用户在购买前了解价格变化情况，也便于后台追踪商品价格调整过程，如表4.12所示。", style_name),
        ("p", "表4.12  价格历史表", caption_style),
        ("t", [
            header,
            ["id", "bigint", "20", "记录ID", "是", "-"],
            ["product_id", "bigint", "20", "商品ID", "否", "-"],
            ["price", "decimal", "10,2", "记录时价格", "否", "-"],
            ["original_price", "decimal", "10,2", "原价", "否", "NULL"],
            ["recorded_time", "datetime", "-", "记录时间", "否", "CURRENT_TIMESTAMP"],
            ["change_type", "varchar", "20", "变化类型", "否", "NULL"],
            ["change_amount", "decimal", "10,2", "变化金额", "否", "NULL"],
            ["change_rate", "decimal", "5,2", "变化比例", "否", "NULL"],
        ], table_style),
        ("p", "（13）降价提醒表用于保存用户针对某个商品设置的目标价格和提醒状态。该表包括提醒ID、用户ID、商品ID、目标价格、设置时当前价格、提醒状态、触发时间、触发价格、通知标识和创建时间等字段。当商品价格低于或等于目标价格时，系统会将提醒状态更新为已触发，并生成相应通知，帮助用户及时了解价格变化，如表4.13所示。", style_name),
        ("p", "表4.13  降价提醒表", caption_style),
        ("t", [
            header,
            ["id", "bigint", "20", "提醒ID", "是", "-"],
            ["user_id", "bigint", "20", "用户ID", "否", "-"],
            ["product_id", "bigint", "20", "商品ID", "否", "-"],
            ["target_price", "decimal", "10,2", "目标价格", "否", "-"],
            ["current_price", "decimal", "10,2", "设置时当前价格", "否", "-"],
            ["status", "tinyint", "4", "状态", "否", "0"],
            ["triggered_time", "datetime", "-", "触发时间", "否", "NULL"],
            ["triggered_price", "decimal", "10,2", "触发价格", "否", "NULL"],
            ["notified", "tinyint", "4", "是否已通知", "否", "0"],
            ["created_time", "datetime", "-", "创建时间", "否", "CURRENT_TIMESTAMP"],
            ["updated_time", "datetime", "-", "更新时间", "否", "CURRENT_TIMESTAMP"],
        ], table_style),
        ("p", "（14）消费预算表用于记录用户每个月设置的预算金额和预警阈值，是理性消费助手中预算进度、剩余金额和超支提醒的基础数据。该表包括预算ID、用户ID、月度预算金额、预算年月、是否启用提醒、预警阈值、创建时间和更新时间等字段。系统按照用户和月份保持唯一预算记录，并结合已支付订单金额计算预算使用情况，如表4.14所示。", style_name),
        ("p", "表4.14  消费预算表", caption_style),
        ("t", [
            header,
            ["id", "bigint", "20", "预算ID", "是", "-"],
            ["user_id", "bigint", "20", "用户ID", "否", "-"],
            ["monthly_budget", "decimal", "10,2", "月度预算金额", "否", "-"],
            ["budget_month", "varchar", "6", "预算年月", "否", "-"],
            ["alert_enabled", "tinyint", "4", "是否启用预算提醒", "否", "1"],
            ["alert_threshold", "int", "11", "预算警告阈值", "否", "80"],
            ["created_time", "datetime", "-", "创建时间", "否", "CURRENT_TIMESTAMP"],
            ["updated_time", "datetime", "-", "更新时间", "否", "CURRENT_TIMESTAMP"],
        ], table_style),
    ]

    current = anchor
    for block in blocks:
        if block[0] == "p":
            current = insert_paragraph_after(doc, current, block[1], block[2])
        else:
            current = insert_table_after(doc, current, block[1], block[2])


def update_test_table(doc: Document) -> None:
    for table in doc.tables:
        if len(table.rows) >= 4 and "用户申请取消" in table.cell(3, 1).text:
            table.cell(2, 2).text = "支付状态由“未支付”变为“已支付”，订单状态进入“待发货”"
            table.cell(2, 3).text = "支付状态由“未支付”变为“已支付”，订单状态进入“待发货”"
            table.cell(3, 2).text = "状态先变为“申请取消中”，再变为“已取消”且库存回退"
            table.cell(3, 3).text = "状态先变为“申请取消中”，再变为“已取消”且库存回退"
            break


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)

    replace_paragraphs(doc)
    update_use_case_tables(doc)
    update_physical_tables(doc)
    append_new_physical_tables(doc)
    update_test_table(doc)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
