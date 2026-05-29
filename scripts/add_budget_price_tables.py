from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


SOURCE = Path(
    r"C:\Users\Administrator\Desktop\我的毕业设计材料\毕业设计终极版\初稿终极版\真终极版压缩包\终极修复版\2203010614张津滔毕业设计定稿.docx"
)
OUT_DIR = Path(r"D:\graduation-project\docs\thesis-alignment")
OUTPUT = OUT_DIR / "2203010614张津滔毕业设计定稿-补预算价格历史表版.docx"


def iter_blocks(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc), "p"
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc), "t"


def normalize(text: str) -> str:
    return text.replace("\xa0", " ").strip()


def find_caption_table(doc: Document, caption: str) -> Table:
    seen = False
    for block, kind in iter_blocks(doc):
        if kind == "p" and caption in normalize(block.text):
            seen = True
            continue
        if seen and kind == "t":
            return block
    raise RuntimeError(f"Cannot find table after caption: {caption}")


def insert_paragraph_after(doc: Document, after_element, text: str, style_name: str):
    paragraph = doc.add_paragraph(text)
    try:
        paragraph.style = style_name
    except KeyError:
        pass
    after_element.addnext(paragraph._p)
    return paragraph._p


def insert_table_after(doc: Document, after_element, rows: list[list[str]], style):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    if style is not None:
        table.style = style
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value
    after_element.addnext(table._tbl)
    return table._tbl


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, OUTPUT)

    doc = Document(OUTPUT)
    table_411 = find_caption_table(doc, "表4.11")
    anchor = table_411._tbl
    table_style = table_411.style
    body_style = "Body Text"
    caption_style = "Caption"
    header = ["字段名称", "类型", "长度", "字段说明", "主键", "默认值"]

    additions = [
        (
            "（12）价格历史表用于保存商品价格随时间变化的记录，是商品详情页展示价格走势、最低价、最高价和平均价格的基础数据。该表包括价格历史ID、商品ID、记录价格、原价、记录时间、变化类型、变化金额和变化比例等字段。商品新增或价格发生调整时，系统会写入对应的价格记录，方便用户在购买前了解商品价格变化情况，也便于后台追踪商品定价调整过程，如表4.12所示。",
            "表4.12  价格历史表",
            [
                header,
                ["id", "bigint", "20", "记录ID", "是", "-"],
                ["product_id", "bigint", "20", "商品ID", "否", "-"],
                ["price", "decimal", "10,2", "记录时价格", "否", "-"],
                ["original_price", "decimal", "10,2", "原价", "否", "NULL"],
                ["recorded_time", "datetime", "-", "记录时间", "否", "CURRENT_TIMESTAMP"],
                ["change_type", "varchar", "20", "变化类型", "否", "NULL"],
                ["change_amount", "decimal", "10,2", "变化金额", "否", "NULL"],
                ["change_rate", "decimal", "5,2", "变化比例", "否", "NULL"],
            ],
        ),
        (
            "（13）消费预算表用于记录用户每个月设置的预算金额和预警阈值，是理性消费助手中预算进度、剩余金额和超支提醒的基础数据。该表包括预算ID、用户ID、月度预算金额、预算年月、是否启用预算提醒、预警阈值、创建时间和更新时间等字段。系统按照用户和月份保存预算记录，并结合已支付订单金额计算当月预算使用情况，如表4.13所示。",
            "表4.13  消费预算表",
            [
                header,
                ["id", "bigint", "20", "预算ID", "是", "-"],
                ["user_id", "bigint", "20", "用户ID", "否", "-"],
                ["monthly_budget", "decimal", "10,2", "月度预算金额", "否", "-"],
                ["budget_month", "varchar", "6", "预算年月", "否", "-"],
                ["alert_enabled", "tinyint", "4", "是否启用预算提醒", "否", "1"],
                ["alert_threshold", "int", "11", "预算警告阈值", "否", "80"],
                ["created_time", "datetime", "-", "创建时间", "否", "CURRENT_TIMESTAMP"],
                ["updated_time", "datetime", "-", "更新时间", "否", "CURRENT_TIMESTAMP"],
            ],
        ),
    ]

    current = anchor
    for description, caption, rows in additions:
        current = insert_paragraph_after(doc, current, description, body_style)
        current = insert_paragraph_after(doc, current, caption, caption_style)
        current = insert_table_after(doc, current, rows, table_style)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
